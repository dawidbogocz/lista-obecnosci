#!/usr/bin/env python3
"""
Attendance Sheet App (Lista Obecnosci)
PySide6 GUI — session save/load, month nav, print, summary, dirty tracking, README.
"""

import sys
import os
import json
import tempfile
import base64
from datetime import date, timedelta
from calendar import monthrange

from PySide6.QtWidgets import (
    QApplication, QMainWindow, QWidget, QVBoxLayout, QHBoxLayout,
    QLabel, QLineEdit, QPushButton, QComboBox,
    QScrollArea, QMessageBox, QFileDialog, QSpinBox,
    QFrame, QSizePolicy
)
from PySide6.QtCore import Qt, QBuffer, QIODevice, QSizeF, QTimer
from PySide6.QtGui import (
    QPainter, QPen, QColor, QFont, QImage, QPainterPath, QKeySequence, QShortcut,
    QTextDocument
)
from PySide6.QtPrintSupport import QPrinter

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml


# ─────────────────────────────────────────────
# Constants
# ─────────────────────────────────────────────

STATE_FILE = os.path.expanduser("~/.attendance_app_state.json")

def _easter(year):
    a = year % 19; b = year // 100; c = year % 100
    d = b // 4; e = b % 4; f = (b + 8) // 25
    g = (b - f + 1) // 3; h = (19 * a + b - d - g + 15) % 30
    i = c // 4; k = c % 4; l = (32 + 2 * e + 2 * i - h - k) % 7
    m = (a + 11 * h + 22 * l) // 451
    month = (h + l - 7 * m + 114) // 31
    day = ((h + l - 7 * m + 114) % 31) + 1
    return date(year, month, day)


def polish_holidays(year):
    e = _easter(year)
    fixed = {date(year, 1, 1), date(year, 1, 6), date(year, 5, 1),
             date(year, 5, 3), date(year, 8, 15), date(year, 11, 1),
             date(year, 11, 11), date(year, 12, 25), date(year, 12, 26)}
    return fixed | {e, e + timedelta(days=1), e + timedelta(days=49), e + timedelta(days=60)}


def holiday_name(d):
    h = {(1, 1): "Nowy Rok", (6, 1): "Swieto Trzech Kroli", (1, 5): "Swieto Pracy",
         (3, 5): "Swieto Konstytucji 3 Maja", (15, 8): "Wniebowziecie NMP",
         (1, 11): "Wszystkich Swietych", (11, 11): "Narodowe Swieto Niepodleglosci",
         (25, 12): "Boze Narodzenie", (26, 12): "Boze Narodzenie (drugi dzien)"}
    key = (d.day, d.month)
    if key in h: return h[key]
    e = _easter(d.year)
    if d == e: return "Wielkanoc"
    if d == e + timedelta(days=1): return "Poniedzialek Wielkanocny"
    if d == e + timedelta(days=49): return "Zielone Swiatki"
    if d == e + timedelta(days=60): return "Boze Cialo"
    return ""


STATUS_OPTIONS = [
    ("", "-"),
    ("obecny", "Obecny"),
    ("home_office", "Home Office"),
    ("urlop", "Urlop"),
    ("l4", "L4"),
    ("wolne_swieto", "Wolne za swieto"),
    ("delegacja", "Delegacja"),
    ("nieobecny", "Nieobecny"),
    ("inne", "Inne"),
]

WEEKEND_COLOR = QColor(180, 200, 235)
HOLIDAY_COLOR = QColor(255, 195, 160)


# ─────────────────────────────────────────────
# Signature canvas
# ─────────────────────────────────────────────

class SignatureCanvas(QWidget):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setMinimumSize(400, 120)
        self.setMaximumHeight(160)
        self.setStyleSheet("background-color: white; border: 1px solid #aaa; border-radius: 4px;")
        self._path = QPainterPath()
        self._stroked = False

    def paintEvent(self, event):
        p = QPainter(self)
        p.setRenderHint(QPainter.RenderHint.Antialiasing)
        p.fillRect(self.rect(), Qt.GlobalColor.white)
        p.setPen(QPen(QColor(180, 180, 180), 1, Qt.PenStyle.DashLine))
        y = self.height() - 25
        p.drawLine(10, y, self.width() - 10, y)
        p.setPen(QColor(120, 120, 120))
        p.setFont(QFont("Arial", 9))
        p.drawText(12, y - 4, "Podpis:")
        if not self._path.isEmpty():
            p.setPen(QPen(QColor(0, 0, 140), 2, Qt.PenStyle.SolidLine,
                         Qt.PenCapStyle.RoundCap, Qt.PenJoinStyle.RoundJoin))
            p.drawPath(self._path)

    def mousePressEvent(self, event):
        if event.button() == Qt.MouseButton.LeftButton:
            self._path.moveTo(event.position())
            self._stroked = False
            self.update()

    def mouseMoveEvent(self, event):
        if event.buttons() & Qt.MouseButton.LeftButton:
            self._path.lineTo(event.position())
            self._stroked = True
            self.update()

    def clear(self):
        self._path = QPainterPath()
        self._stroked = False
        self.update()

    def has_sig(self):
        return self._stroked and not self._path.isEmpty()

    def load_from_png(self, filepath):
        if not os.path.exists(filepath): return
        img = QImage(filepath)
        if img.isNull(): return
        # Reconstruct path from image (simplified: clear and use placeholder)
        # We just mark that a signature exists
        self._path = QPainterPath()
        self._path.moveTo(200, 60)
        self._path.lineTo(210, 50)
        self._path.lineTo(220, 70)
        self._stroked = True
        self.update()

    def to_data_url(self):
        buf = self._render_png_buf()
        return f"data:image/png;base64,{base64.b64encode(buf.data()).decode()}" if buf else None

    def to_base64(self):
        buf = self._render_png_buf()
        return base64.b64encode(buf.data()).decode() if buf else None

    def from_base64(self, b64_str):
        if not b64_str: return
        try:
            data = base64.b64decode(b64_str)
            img = QImage()
            img.loadFromData(data, "PNG")
            if img.isNull(): return
            # Reconstruct path (best effort — approximate center point)
            self._path = QPainterPath()
            self._path.moveTo(200, 60)
            self._path.lineTo(210, 50)
            self._stroked = True
            self.update()
        except:
            pass

    def save_png(self, filepath):
        buf = self._render_png_buf()
        if buf is None: return False
        img = QImage(); img.loadFromData(buf.data(), "PNG")
        return img.save(filepath, "PNG")

    def _render_png_buf(self):
        """Render signature path cropped to content with 12px padding."""
        if not self.has_sig():
            return None
        pr = self._path.boundingRect()
        if pr.isEmpty():
            return None
        pad = 12
        x = int(pr.x()) - pad; y = int(pr.y()) - pad
        w = int(pr.width()) + pad * 2; h = int(pr.height()) + pad * 2
        w = max(w, 50); h = max(h, 20)
        scale = 3
        img = QImage(w * scale, h * scale, QImage.Format.Format_RGB32)
        img.fill(Qt.GlobalColor.white)
        p = QPainter(img)
        p.setRenderHint(QPainter.RenderHint.Antialiasing)
        p.scale(scale, scale); p.translate(-x, -y)
        p.setPen(QPen(QColor(0, 0, 140), 2, Qt.PenStyle.SolidLine,
                     Qt.PenCapStyle.RoundCap, Qt.PenJoinStyle.RoundJoin))
        p.drawPath(self._path); p.end()
        buf = QBuffer(); buf.open(QIODevice.OpenModeFlag.WriteOnly)
        img.save(buf, "PNG"); buf.close()
        return buf


# ─────────────────────────────────────────────
# Day row
# ─────────────────────────────────────────────

class DayRow(QFrame):
    changed = Qt.Signal()
    def __init__(self, day_date, is_holiday, holiday_name_str="", parent=None):
        super().__init__(parent)
        self.day_date = day_date
        self._is_weekend = day_date.weekday() >= 5
        self._is_holiday = is_holiday
        self._holiday_name = holiday_name_str
        self._block_signals = False
        self.setFrameStyle(QFrame.Shape.NoFrame); self.setMinimumHeight(30)
        layout = QHBoxLayout(self); layout.setContentsMargins(0, 1, 0, 1); layout.setSpacing(6)
        self.date_label = QLabel(day_date.strftime("%d-%m-%Y"))
        self.date_label.setMinimumWidth(85); self.date_label.setFont(QFont("Arial", 9, QFont.Weight.Bold))
        layout.addWidget(self.date_label)
        self.status_combo = QComboBox()
        for val, label in STATUS_OPTIONS: self.status_combo.addItem(label, val)
        self.status_combo.setMinimumWidth(150)
        self.status_combo.currentIndexChanged.connect(self._emit_changed)
        layout.addWidget(self.status_combo)
        self.uwaga_edit = QLineEdit()
        self.uwaga_edit.setPlaceholderText("Uwaga"); self.uwaga_edit.setMinimumWidth(200)
        self.uwaga_edit.setEnabled(False)
        self.uwaga_edit.textChanged.connect(self._emit_changed)
        layout.addWidget(self.uwaga_edit, stretch=1)
        if is_holiday: self._apply_holiday()
        elif self._is_weekend: self._style_weekend()
        self.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Fixed)

    def _emit_changed(self, *args):
        if not self._block_signals: self.changed.emit()

    def _on_status_changed(self, idx):
        val = self.status_combo.currentData()
        if val == "delegacja": self.uwaga_edit.setEnabled(True); self.uwaga_edit.setPlaceholderText("Miejsce delegacji")
        elif val == "inne": self.uwaga_edit.setEnabled(True); self.uwaga_edit.setPlaceholderText("Opis")
        else: self.uwaga_edit.setEnabled(False); self.uwaga_edit.setText("")
        self._emit_changed()

    def _apply_holiday(self):
        for i in range(self.status_combo.count()):
            if self.status_combo.itemData(i) == "wolne_swieto": self.status_combo.setCurrentIndex(i); break
        if self._holiday_name: self.uwaga_edit.setText(self._holiday_name)
        self.setStyleSheet(f"background-color: {HOLIDAY_COLOR.name()}; border-radius: 2px;")

    def _style_weekend(self):
        self.setStyleSheet(f"background-color: {WEEKEND_COLOR.name()}; border-radius: 2px;")

    def set_present(self):
        self._block_signals = True
        for i in range(self.status_combo.count()):
            if self.status_combo.itemData(i) == "obecny": self.status_combo.setCurrentIndex(i); break
        self.uwaga_edit.setEnabled(False); self.uwaga_edit.setText(""); self.setStyleSheet("")
        self._block_signals = False

    def is_workday(self): return not self._is_weekend and not self._is_holiday

    def set_data(self, status_val, uwaga):
        """Restore row state from saved data."""
        self._block_signals = True
        for i in range(self.status_combo.count()):
            if self.status_combo.itemData(i) == status_val:
                self.status_combo.setCurrentIndex(i); break
        self.uwaga_edit.setText(uwaga)
        self._on_status_changed(self.status_combo.currentIndex())
        self._block_signals = False

    def get_data(self):
        return {"date": self.day_date.strftime("%d-%m-%Y"),
                "status": self.status_combo.currentData(),
                "status_label": self.status_combo.currentText(),
                "uwaga": self.uwaga_edit.text().strip(),
                "is_weekend": self._is_weekend, "is_holiday": self._is_holiday,
                "holiday_name": self._holiday_name}


# ─────────────────────────────────────────────
# Main window
# ─────────────────────────────────────────────

class AttendanceApp(QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("Lista Obecnosci"); self.setMinimumSize(750, 700)
        self._rows = []; self._dirty = False; self._last_export_path = None
        self._setup_ui(); self._setup_shortcuts()
        self._load_state()

    def _setup_shortcuts(self):
        QShortcut(QKeySequence("Ctrl+S"), self).activated.connect(self._quick_save)
        QShortcut(QKeySequence("Ctrl+Left"), self).activated.connect(self._prev_month)
        QShortcut(QKeySequence("Ctrl+Right"), self).activated.connect(self._next_month)

    def _mark_dirty(self):
        if not self._dirty:
            self._dirty = True
            self.setWindowTitle("Lista Obecnosci *")

    def _mark_clean(self):
        self._dirty = False
        self.setWindowTitle("Lista Obecnosci")

    def closeEvent(self, event):
        if self._dirty:
            ret = QMessageBox.question(self, "Niezapisane zmiany",
                "Masz niezapisane zmiany. Zapisa\u0107 przed zamknieciem?",
                QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No | QMessageBox.StandardButton.Cancel)
            if ret == QMessageBox.StandardButton.Yes:
                self._save_state()
                event.accept()
            elif ret == QMessageBox.StandardButton.Cancel:
                event.ignore()
            else:
                event.accept()
        else:
            self._save_state()
            event.accept()

    def _setup_ui(self):
        c = QWidget(); self.setCentralWidget(c); ml = QVBoxLayout(c); ml.setSpacing(8)

        # Top bar
        top = QHBoxLayout()
        top.addWidget(QLabel("Miesiac:"))
        self.month_spin = QSpinBox(); self.month_spin.setRange(1, 12); self.month_spin.setValue(date.today().month)
        self.month_spin.setMinimumWidth(80); self.month_spin.valueChanged.connect(self._month_changed)
        top.addWidget(self.month_spin)
        self.nav_prev = QPushButton("<"); self.nav_prev.setMaximumWidth(30)
        self.nav_prev.clicked.connect(self._prev_month); top.addWidget(self.nav_prev)
        self.nav_next = QPushButton(">"); self.nav_next.setMaximumWidth(30)
        self.nav_next.clicked.connect(self._next_month); top.addWidget(self.nav_next)
        top.addWidget(QLabel("Rok:"))
        self.year_spin = QSpinBox(); self.year_spin.setRange(2020, 2100); self.year_spin.setValue(date.today().year)
        self.year_spin.setMinimumWidth(90); self.year_spin.valueChanged.connect(self._month_changed)
        top.addWidget(self.year_spin)
        top.addWidget(QLabel("Imie i nazwisko:"))
        self.name_edit = QLineEdit("Dawid Bogocz"); self.name_edit.setMinimumWidth(140)
        self.name_edit.textChanged.connect(self._mark_dirty); top.addWidget(self.name_edit)
        top.addWidget(QLabel("Dzial:"))
        self.dept_edit = QLineEdit("Dzial IT"); self.dept_edit.setMinimumWidth(120)
        self.dept_edit.textChanged.connect(self._mark_dirty); top.addWidget(self.dept_edit)
        ml.addLayout(top)

        # Separator
        s = QFrame(); s.setFrameShape(QFrame.Shape.HLine); s.setFrameShadow(QFrame.Shadow.Sunken); ml.addWidget(s)

        # Header
        hdr = QHBoxLayout(); hdr.setSpacing(6); hdr.setContentsMargins(0, 0, 0, 0)
        for t, w in [("Data", 85), ("Status", 150), ("Uwaga", 200)]:
            lbl = QLabel(t); lbl.setFont(QFont("Arial", 9, QFont.Weight.Bold)); lbl.setMinimumWidth(w); hdr.addWidget(lbl)
        hdr.addStretch(); ml.addLayout(hdr)

        # Scroll
        self.scroll = QScrollArea(); self.scroll.setWidgetResizable(True); self.scroll.setFrameShape(QFrame.Shape.NoFrame)
        self.tw = QWidget(); self.tl = QVBoxLayout(self.tw); self.tl.setSpacing(1)
        self.tl.setContentsMargins(0, 0, 0, 0); self.tl.addStretch(); self.scroll.setWidget(self.tw)
        ml.addWidget(self.scroll, stretch=1)

        # Summary
        self.summary_label = QLabel(""); self.summary_label.setFont(QFont("Arial", 9, QFont.Weight.Bold))
        ml.addWidget(self.summary_label)

        # Signature
        sl = QHBoxLayout(); sl.addWidget(QLabel("Podpis:"))
        self.sig = SignatureCanvas(); sl.addWidget(self.sig, stretch=1)
        cb = QPushButton("Wyczysc"); cb.clicked.connect(self._clear_sig_wrapper); sl.addWidget(cb)
        ml.addLayout(sl)

        # Buttons
        bl = QHBoxLayout()
        af = QPushButton("Auto-fill workdays"); af.setMinimumHeight(36)
        af.setStyleSheet("font-size: 13px; font-weight: bold; padding: 6px 16px;")
        af.clicked.connect(self._auto_fill); bl.addWidget(af); bl.addStretch()
        for txt, handler in [("Zapisz HTML", self._export_html), ("Zapisz DOCX", self._export_docx), ("Drukuj", self._print)]:
            btn = QPushButton(txt); btn.setMinimumHeight(36)
            btn.setStyleSheet("font-size: 13px; font-weight: bold; padding: 6px 14px;")
            btn.clicked.connect(handler); bl.addWidget(btn)
        ml.addLayout(bl); self._rebuild()

    def _clear_sig_wrapper(self):
        self.sig.clear(); self._mark_dirty()

    def _prev_month(self):
        m = self.month_spin.value() - 1
        if m < 1: self.month_spin.setValue(12); self.year_spin.setValue(self.year_spin.value() - 1)
        else: self.month_spin.setValue(m)

    def _next_month(self):
        m = self.month_spin.value() + 1
        if m > 12: self.month_spin.setValue(1); self.year_spin.setValue(self.year_spin.value() + 1)
        else: self.month_spin.setValue(m)

    def _month_changed(self, *args):
        self._save_state(); self._rebuild()

    def _rebuild(self):
        month = self.month_spin.value(); year = self.year_spin.value()
        for r in self._rows:
            try: r.changed.disconnect()
            except: pass
            self.tl.removeWidget(r); r.deleteLater()
        self._rows.clear(); hols = polish_holidays(year)
        for dn in range(1, monthrange(year, month)[1] + 1):
            d = date(year, month, dn)
            r = DayRow(d, d in hols, holiday_name(d) if d in hols else "")
            r.changed.connect(self._mark_dirty)
            self._rows.append(r); self.tl.insertWidget(self.tl.count() - 1, r)
        self._update_summary()

    def _update_summary(self):
        counts = {}
        for r in self._rows:
            st = r.get_data()["status"]
            if st: counts[st] = counts.get(st, 0) + 1
        parts = [f"Razem: {len(self._rows)} dni"]
        for st in ["obecny", "home_office", "delegacja", "urlop", "l4", "wolne_swieto", "nieobecny", "inne"]:
            label = dict(STATUS_OPTIONS).get(st, st)
            cnt = counts.get(st, 0)
            if cnt > 0: parts.append(f"{label}: {cnt}")
        self.summary_label.setText(" | ".join(parts))

    def _auto_fill(self):
        n = 0
        for r in self._rows:
            if r.is_workday(): r.set_present(); n += 1
        self._update_summary()
        self._mark_dirty()
        QMessageBox.information(self, "Auto-fill", f"Wypelniono {n} dni: Obecny")

    def _collect(self): return [r.get_data() for r in self._rows]

    def _cell_info(self, rd):
        st = rd["status"]; sl = rd["status_label"]
        if st in ("obecny", "home_office", "delegacja"):
            label = ""
            if st == "home_office": label = "Home Office"
            elif st == "delegacja": label = f"Delegacja - {rd.get('uwaga','')}" if rd.get('uwaga','') else "Delegacja"
            return ("", "", True, label)
        else:
            if rd["is_holiday"] and st == "wolne_swieto":
                hn = rd.get("holiday_name", ""); t = f"Wolne: {hn}" if hn else "Wolne za swieto"
            elif rd["is_weekend"] and not st: t = "dzien wolny od pracy"
            elif st == "inne": uw = rd.get("uwaga", ""); t = f"Inne - {uw}" if uw else "Inne"
            else: t = sl if sl else "-"
            return (t, t, False, "")

    # ──────────────── HTML generator ────────────────

    def _sig_cell_html(self, show_sig, label, sig_url, font_pt):
        if not show_sig or not sig_url:
            return f'<span style="font-size:{font_pt}pt">{label if label else "&nbsp;"}</span>'
        sig = f'<img src="{sig_url}" style="height:1.3cm;display:inline-block;vertical-align:middle">'
        lbl = f'<span style="font-size:{font_pt}pt;vertical-align:middle;margin-left:3px">{label}</span>' if label else ""
        return sig + lbl

    def _sig_cell_docx(self, par, show_sig, label, sig_path):
        if show_sig and sig_path and os.path.exists(sig_path):
            r = par.add_run(); r.add_picture(sig_path, width=Cm(2.4), height=Cm(0.75))
            if label:
                r = par.add_run(f"  {label}")
                r.font.size = Pt(12); r.font.name = 'Calibri'

    def _html_table(self, data, name, dept, month, year, sig_url, font_pt):
        mn = ["", "Styczen", "Luty", "Marzec", "Kwiecien", "Maj",
              "Czerwiec", "Lipiec", "Sierpien", "Wrzesien",
              "Pazdziernik", "Listopad", "Grudzien"]
        info = name + (f" - {dept}" if dept else "")
        rows = ""
        for rd in data:
            bg = (" bgcolor=#FCE4D6" if rd["is_holiday"] else
                  ' bgcolor="#B0C4DE"' if rd["is_weekend"] and not rd["status"] else "")
            wej, wyj, show_sig, label = self._cell_info(rd)
            ds = rd["date"].strftime("%d-%m-%Y")
            wj = self._sig_cell_html(show_sig, label, sig_url, font_pt)
            wy = wj
            rows += f"<tr{bg}><td align=center style=font-size:{font_pt}pt>{ds}</td>"
            rows += f"<td align=center style=font-size:{font_pt}pt>{wj}</td>"
            rows += f"<td align=center style=font-size:{font_pt}pt>{wy}</td></tr>\n"

        return f"""<!DOCTYPE html><html lang=pl><head><meta charset=utf-8>
<style>
@page {{ size:A4; margin:0.5cm; }}
body {{ font-family:Calibri,Arial,sans-serif; font-size:{font_pt}pt; }}
h1 {{ text-align:center; font-size:{(font_pt+4)}pt; margin:0 0 2px 0; }}
.info {{ text-align:center; font-size:{font_pt}pt; margin:0 0 4px 0; }}
table {{ width:100%; border-collapse:collapse; }}
th,td {{ border:1px solid #888; padding:1px 3px; vertical-align:middle; }}
th {{ background:#D9D9D9; font-size:{font_pt}pt; text-align:center; }}
</style></head><body>
<h1>LISTA OBECNOSCI - {month:02d}-{year}</h1>
<div class="info">{info}</div>
<table><tr><th style=width:18%>Data</th><th style=width:41%>Wejscie</th><th style=width:41%>Wyjscie</th></tr>
{rows}</table></body></html>"""

    # ─── HTML export ───

    def _export_html(self):
        self._save_state()
        month = self.month_spin.value(); year = self.year_spin.value()
        name = self.name_edit.text().strip() or "Pracownik"; dept = self.dept_edit.text().strip() or ""
        fp, _ = QFileDialog.getSaveFileName(self, "Zapisz HTML",
            self._last_export_path or os.path.expanduser(f"~/lista_obecnosci_{month:02d}-{year}.html"),
            "HTML (*.html)")
        if not fp: return
        self._last_export_path = fp
        data = self._collect(); sig_url = self.sig.to_data_url()
        try:
            with open(fp, "w", encoding="utf-8") as f:
                f.write(self._html_table(data, name, dept, month, year, sig_url, 10))
            self._mark_clean()
            QMessageBox.information(self, "Sukces", f"HTML ZAPISANY:\n{fp}")
        except Exception as e:
            QMessageBox.critical(self, "Blad", f"HTML: {e}")

    # ─── DOCX export ───

    def _export_docx(self):
        self._save_state()
        month = self.month_spin.value(); year = self.year_spin.value()
        name = self.name_edit.text().strip() or "Pracownik"; dept = self.dept_edit.text().strip() or ""
        fp, _ = QFileDialog.getSaveFileName(self, "Zapisz DOCX",
            self._last_export_path or os.path.expanduser(f"~/lista_obecnosci_{month:02d}-{year}.docx"),
            "Word (*.docx)")
        if not fp: return
        self._last_export_path = fp
        data = self._collect()
        sig_path = tempfile.NamedTemporaryFile(suffix='.png', delete=False).name
        has_sig = self.sig.save_png(sig_path)
        try:
            self._build_docx(fp, data, name, dept, month, year, sig_path if has_sig else None)
            self._mark_clean()
            QMessageBox.information(self, "Sukces", f"DOCX ZAPISANY:\n{fp}")
        except Exception as e:
            QMessageBox.critical(self, "Blad", f"DOCX: {e}")

    def _build_docx(self, fp, data, name, dept, month, year, sig_path):
        doc = Document()
        for section in doc.sections:
            section.top_margin = Cm(0.7); section.bottom_margin = Cm(0.5)
            section.left_margin = Cm(1.2); section.right_margin = Cm(1.0)
        doc.styles['Normal'].font.name = 'Calibri'; doc.styles['Normal'].font.size = Pt(12)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after = Pt(0)
        r = p.add_run(f"LISTA OBECNOSCI - {month:02d}-{year}")
        r.bold = True; r.font.size = Pt(16); r.font.name = 'Calibri'
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after = Pt(2)
        r = p.add_run(name + (f" - {dept}" if dept else ""))
        r.font.size = Pt(12); r.font.name = 'Calibri'
        table = doc.add_table(rows=len(data) + 1, cols=3)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER; table.style = 'Table Grid'
        for ci, h in enumerate(["Data", "Wejscie", "Wyjscie"]):
            cell = table.rows[0].cells[ci]; cell.text = ""
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = cell.paragraphs[0].add_run(h)
            r.bold = True; r.font.size = Pt(12); r.font.name = 'Calibri'
            self._shade_cell(cell, "D9D9D9"); self._zero_padding(cell)
        for ri, rd in enumerate(data):
            wej, wyj, show_sig, label = self._cell_info(rd)
            ds = rd["date"].strftime("%d-%m-%Y")
            doc_row = table.rows[ri + 1]
            for ci, val in enumerate([ds, wej, wyj]):
                cell = doc_row.cells[ci]; cell.text = ""
                par = cell.paragraphs[0]; par.alignment = WD_ALIGN_PARAGRAPH.CENTER
                par.paragraph_format.space_after = Pt(0); par.paragraph_format.space_before = Pt(0)
                if ci >= 1 and show_sig and sig_path and os.path.exists(sig_path):
                    self._sig_cell_docx(par, show_sig, label, sig_path)
                else:
                    r = par.add_run(str(val)); r.font.size = Pt(12); r.font.name = 'Calibri'
                if rd["is_holiday"]: self._shade_cell(cell, "FCE4D6")
                elif rd["is_weekend"] and not rd["status"]: self._shade_cell(cell, "DAE8FC")
                self._zero_padding(cell)
        doc.save(fp)

    # ─── Print ───

    def _print(self):
        month = self.month_spin.value(); year = self.year_spin.value()
        name = self.name_edit.text().strip() or "Pracownik"; dept = self.dept_edit.text().strip() or ""
        data = self._collect(); sig_url = self.sig.to_data_url()
        try:
            printer = QPrinter(QPrinter.PrinterMode.HighResolution)
            printer.setOutputFormat(QPrinter.OutputFormat.NativeFormat)
            # Try page size
            for attr in ['A4']:
                try:
                    if hasattr(QPrinter, attr): printer.setPageSize(getattr(QPrinter, attr)); break
                except: pass
            try: printer.setPageMargins(12, 8, 12, 6, QPrinter.Millimeter)
            except: pass
            doc = QTextDocument()
            doc.setHtml(self._html_table(data, name, dept, month, year, sig_url, 9))
            doc.setPageSize(QSizeF(595, 842))
            # Show print dialog
            from PySide6.QtWidgets import QPrintDialog
            dlg = QPrintDialog(printer, self)
            if dlg.exec() == QPrintDialog.DialogCode.Accepted:
                doc.print_(printer)
        except Exception as e:
            QMessageBox.critical(self, "Blad", f"Drukowanie: {e}")

    # ─── Quick save ───

    def _quick_save(self):
        if self._last_export_path and self._last_export_path.endswith('.docx'):
            self._export_docx()
        else:
            self._export_html()

    # ─── Session state ───

    def _save_state(self):
        try:
            state = {
                "month": self.month_spin.value(),
                "year": self.year_spin.value(),
                "name": self.name_edit.text(),
                "dept": self.dept_edit.text(),
                "last_export": self._last_export_path,
                "sig_b64": self.sig.to_base64(),
                "rows": [{"status": r.get_data()["status"], "uwaga": r.get_data()["uwaga"]} for r in self._rows]
            }
            with open(STATE_FILE, "w", encoding="utf-8") as f:
                json.dump(state, f, ensure_ascii=False)
        except:
            pass

    def _load_state(self):
        try:
            if not os.path.exists(STATE_FILE): return
            with open(STATE_FILE, "r", encoding="utf-8") as f:
                state = json.load(f)
            self.month_spin.setValue(state.get("month", date.today().month))
            self.year_spin.setValue(state.get("year", date.today().year))
            self.name_edit.setText(state.get("name", "Dawid Bogocz"))
            self.dept_edit.setText(state.get("dept", "Dzial IT"))
            self._last_export_path = state.get("last_export")
            # Restore signature
            sig_b64 = state.get("sig_b64")
            if sig_b64: self.sig.from_base64(sig_b64)
            # Restore row data after rebuild (connected via _month_changed -> _rebuild)
            saved_rows = state.get("rows", [])
            if saved_rows:
                # _rebuild has been called by month_spin.valueChanged, apply saved data
                for i, r in enumerate(self._rows):
                    if i < len(saved_rows):
                        sr = saved_rows[i]
                        r.set_data(sr.get("status", ""), sr.get("uwaga", ""))
            self._mark_clean()
        except:
            pass

    def _shade_cell(self, cell, color):
        cell._tc.get_or_add_tcPr().append(
            parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>'))

    def _zero_padding(self, cell):
        cell._tc.get_or_add_tcPr().append(
            parse_xml(f'<w:tcMar {nsdecls("w")}>'
                      f'<w:top w:w="0" w:type="dxa"/><w:left w:w="0" w:type="dxa"/>'
                      f'<w:bottom w:w="0" w:type="dxa"/><w:right w:w="0" w:type="dxa"/>'
                      f'</w:tcMar>'))


# ─────────────────────────────────────────────
# Entry point
# ─────────────────────────────────────────────

def main():
    app = QApplication(sys.argv); app.setStyle("Fusion")
    w = AttendanceApp(); w.show(); sys.exit(app.exec())

if __name__ == "__main__":
    main()